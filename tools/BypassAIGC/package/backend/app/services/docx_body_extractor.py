from __future__ import annotations

import io
import json
import re
from copy import deepcopy
from dataclasses import dataclass
from difflib import SequenceMatcher
from pathlib import Path
from typing import List, Tuple
from uuid import uuid4

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.config import get_exe_dir

NUMERIC_HEADING_RE = re.compile(r"^[1-9]\d*(?:\.\d+)+(?!\d).*$")
SINGLE_NUMERIC_HEADING_RE = re.compile(r"^\d+\.(?!\d)\s*.*$")
MEASUREMENT_PREFIX_RE = re.compile(
    r"^\d+(?:\.\d+)+(?!\d)\s*"
    r"(?:英寸|寸|毫米|厘米|米|mm|cm|m|in(?:ch(?:es)?)?|px|像素|%|℃|°c)"
    r"(?:$|[\s,，.。;；、的×xX])",
    re.IGNORECASE,
)
COMPACT_NUMERIC_CHAPTER_HEADING_RE = re.compile(r"^\d{1,2}(?![\d\.])[\u4e00-\u9fff].*$")
CHINESE_SECTION_HEADING_RE = re.compile(r"^[一二三四五六七八九十百千]+、\s*.*$")
PAREN_CHINESE_HEADING_RE = re.compile(r"^[（(][一二三四五六七八九十百千]+[）)]\s*.*$")
START_HEADING_RE = re.compile(r"^1\.1(?!\d).*$")
CHAPTER_HEADING_RE = re.compile(r"^第[一二三四五六七八九十百零\d]+章(?:[\s\u3000].*)?$")
TOC_HEADING_RE = re.compile(r"^(目\s*录|目录|table\s+of\s+contents)$", re.IGNORECASE)
ABSTRACT_HEADING_RE = re.compile(r"^(摘\s*要|abstract)(?:[\s:：].*)?$", re.IGNORECASE)
KEYWORDS_HEADING_RE = re.compile(r"^(关\s*键\s*词|key\s*words?)(?:[\s:：].*)?$", re.IGNORECASE)
FRONT_MATTER_HEADING_RE = re.compile(
    r"^(摘\s*要|abstract|关\s*键\s*词|key\s*words?)(?:[\s:：].*)?$",
    re.IGNORECASE,
)
INTRODUCTION_HEADING_RE = re.compile(r"^(前言|序言|绪论|引言)(?:[\s:：].*)?$", re.IGNORECASE)
STOP_SECTION_RE = re.compile(
    r"^(参考文献|references?|致谢|谢辞|附录|appendix(?:es)?)(?:[\s:：].*)?$",
    re.IGNORECASE,
)
SUMMARY_HEADING_RE = re.compile(r"^(总\s*结|结\s*论)(?:[\s:：].*)?$", re.IGNORECASE)
CAPTION_RE = re.compile(
    r"^((图|表)\s*[\d一二三四五六七八九十]+([\-\.．]\d+)*|"
    r"(figure|table)\s*[\d一二三四五六七八九十]+([\-\.．]\d+)*)",
    re.IGNORECASE,
)
PAGE_NUMBER_RE = re.compile(r"^\d+$")
PALE_YELLOW_FILL = "FFF2CC"
WORD_SOURCE_STORE_DIR = Path(get_exe_dir()) / ".word_source_cache"


@dataclass
class ExtractedBodyParagraph:
    source_index: int
    paragraph_type: str
    text: str


@dataclass
class DocxBodyExtractionResult:
    source_filename: str
    output_filename: str
    marked_output_filename: str
    extraction_token: str
    extracted_text: str
    paragraphs: List[ExtractedBodyParagraph]
    start_heading: str


@dataclass
class MarkedDocxResult:
    source_filename: str
    output_filename: str
    start_heading: str
    paragraph_count: int
    content: bytes


@dataclass
class StoredWordSource:
    token: str
    source_filename: str
    start_heading: str
    paragraphs: List[ExtractedBodyParagraph]
    docx_bytes: bytes


@dataclass
class ResultDocxExport:
    source_filename: str
    output_filename: str
    content: bytes


@dataclass
class AigcMatchedDocxExtractionResult:
    docx_result: DocxBodyExtractionResult
    matched_report_segment_count: int
    unmatched_report_segment_count: int


def extract_body_from_word(filename: str, content: bytes) -> DocxBodyExtractionResult:
    """Extract thesis body paragraphs from a Word file."""
    document, normalized_docx_bytes = _load_word_source(filename, content)
    safe_name = filename or "uploaded.docx"
    stem = Path(safe_name).stem

    paragraphs, start_heading = _extract_body_paragraphs(document)
    if not paragraphs:
        raise ValueError("未识别到正文段落，请确认文档中存在 1.1 等正文结构")

    stored = _persist_word_source(safe_name, normalized_docx_bytes, paragraphs, start_heading)
    extracted_text = "\n\n".join(item.text for item in paragraphs)
    return DocxBodyExtractionResult(
        source_filename=safe_name,
        output_filename=f"{stem}_正文提取结果.txt",
        marked_output_filename=f"{stem}_正文提取标记版.docx",
        extraction_token=stored.token,
        extracted_text=extracted_text,
        paragraphs=paragraphs,
        start_heading=start_heading,
    )


def extract_aigc_matched_body_from_word(
    filename: str,
    content: bytes,
    report_segments: List[object],
) -> AigcMatchedDocxExtractionResult:
    """Extract Word body paragraphs matched by AIGC suspicious PDF segments."""
    document, normalized_docx_bytes = _load_word_source(filename, content)
    safe_name = filename or "uploaded.docx"
    stem = Path(safe_name).stem

    paragraphs, _ = _extract_body_paragraphs(document)
    if not paragraphs:
        raise ValueError("未识别到正文段落，请确认文档中存在 1.1 等正文结构")
    if not report_segments:
        raise ValueError("检测报告中未识别到可匹配的疑似片段")

    matched_paragraphs, matched_count = _match_report_segments_to_word_paragraphs(
        paragraphs,
        report_segments,
    )
    if not matched_paragraphs:
        raise ValueError("检测报告中的疑似片段未能匹配到 Word 正文，请确认 Word 与检测报告对应")

    start_heading = "AIGC疑似片段"
    stored = _persist_word_source(
        safe_name,
        normalized_docx_bytes,
        matched_paragraphs,
        start_heading,
    )
    extracted_text = "\n\n".join(item.text for item in matched_paragraphs)
    return AigcMatchedDocxExtractionResult(
        docx_result=DocxBodyExtractionResult(
            source_filename=safe_name,
            output_filename=f"{stem}_疑似片段提取结果.txt",
            marked_output_filename=f"{stem}_疑似片段标记版.docx",
            extraction_token=stored.token,
            extracted_text=extracted_text,
            paragraphs=matched_paragraphs,
            start_heading=start_heading,
        ),
        matched_report_segment_count=matched_count,
        unmatched_report_segment_count=max(0, len(report_segments) - matched_count),
    )


def build_marked_word_from_word(filename: str, content: bytes) -> MarkedDocxResult:
    """Generate a marked DOCX highlighting extracted body paragraphs."""
    document, _ = _load_word_source(filename, content)
    safe_name = filename or "uploaded.docx"
    stem = Path(safe_name).stem

    paragraphs, start_heading = _extract_body_paragraphs(document)
    if not paragraphs:
        raise ValueError("未识别到正文段落，请确认文档中存在 1.1 等正文结构")

    for item in paragraphs:
        _apply_pale_yellow_highlight(document.paragraphs[item.source_index])

    buffer = io.BytesIO()
    document.save(buffer)
    return MarkedDocxResult(
        source_filename=safe_name,
        output_filename=f"{stem}_正文提取标记版.docx",
        start_heading=start_heading,
        paragraph_count=len(paragraphs),
        content=buffer.getvalue(),
    )


def build_marked_word_from_token(extraction_token: str) -> MarkedDocxResult:
    """Generate a marked DOCX highlighting cached extracted Word paragraphs."""
    stored = load_stored_word_source(extraction_token)
    document = Document(io.BytesIO(stored.docx_bytes))

    for item in stored.paragraphs:
        if 0 <= item.source_index < len(document.paragraphs):
            _apply_pale_yellow_highlight(document.paragraphs[item.source_index])

    buffer = io.BytesIO()
    document.save(buffer)
    stem = Path(stored.source_filename).stem
    output_suffix = "疑似片段标记版" if "疑似片段" in stored.start_heading else "正文提取标记版"
    return MarkedDocxResult(
        source_filename=stored.source_filename,
        output_filename=f"{stem}_{output_suffix}.docx",
        start_heading=stored.start_heading,
        paragraph_count=len(stored.paragraphs),
        content=buffer.getvalue(),
    )


def load_stored_word_source(token: str) -> StoredWordSource:
    store_dir = _get_source_dir(token)
    metadata_path = store_dir / "metadata.json"
    docx_path = store_dir / "source.docx"

    if not metadata_path.exists() or not docx_path.exists():
        raise ValueError("未找到对应的 Word 提取缓存，请重新上传并提取正文")

    metadata = json.loads(metadata_path.read_text(encoding="utf-8"))
    paragraphs = [
        ExtractedBodyParagraph(
            source_index=item["source_index"],
            paragraph_type=item.get("paragraph_type", "body"),
            text=item["text"],
        )
        for item in metadata.get("paragraphs", [])
    ]

    return StoredWordSource(
        token=token,
        source_filename=metadata.get("source_filename") or "uploaded.docx",
        start_heading=metadata.get("start_heading") or "",
        paragraphs=paragraphs,
        docx_bytes=docx_path.read_bytes(),
    )


def build_session_paragraphs_from_text(extraction_token: str, text: str) -> List[ExtractedBodyParagraph]:
    stored = load_stored_word_source(extraction_token)
    current_paragraphs = _split_text_into_body_paragraphs(text)

    if len(current_paragraphs) != len(stored.paragraphs):
        raise ValueError(
            "当前文本段落数与提取结果不一致，请保持段落结构一致，或重新上传 Word 提取正文后再开始优化"
        )

    return [
        ExtractedBodyParagraph(
            source_index=stored.paragraphs[index].source_index,
            paragraph_type=stored.paragraphs[index].paragraph_type,
            text=current_paragraphs[index],
        )
        for index in range(len(current_paragraphs))
    ]


def build_result_word_from_source(extraction_token: str, final_paragraphs: List[Tuple[int, str]]) -> ResultDocxExport:
    stored = load_stored_word_source(extraction_token)
    document = Document(io.BytesIO(stored.docx_bytes))

    for source_index, text in final_paragraphs:
        if source_index < 0 or source_index >= len(document.paragraphs):
            continue
        _replace_paragraph_text_preserving_style(document.paragraphs[source_index], text)

    buffer = io.BytesIO()
    document.save(buffer)
    stem = Path(stored.source_filename).stem
    return ResultDocxExport(
        source_filename=stored.source_filename,
        output_filename=f"{stem}_优化结果.docx",
        content=buffer.getvalue(),
    )


def _persist_word_source(
    filename: str,
    docx_bytes: bytes,
    paragraphs: List[ExtractedBodyParagraph],
    start_heading: str,
) -> StoredWordSource:
    token = uuid4().hex
    store_dir = _get_source_dir(token)
    store_dir.mkdir(parents=True, exist_ok=True)

    (store_dir / "source.docx").write_bytes(docx_bytes)
    metadata = {
        "source_filename": filename,
        "start_heading": start_heading,
        "paragraphs": [
            {
                "source_index": item.source_index,
                "paragraph_type": item.paragraph_type,
                "text": item.text,
            }
            for item in paragraphs
        ],
    }
    (store_dir / "metadata.json").write_text(
        json.dumps(metadata, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

    return StoredWordSource(
        token=token,
        source_filename=filename,
        start_heading=start_heading,
        paragraphs=paragraphs,
        docx_bytes=docx_bytes,
    )


def _match_report_segments_to_word_paragraphs(
    paragraphs: List[ExtractedBodyParagraph],
    report_segments: List[object],
) -> Tuple[List[ExtractedBodyParagraph], int]:
    normalized_paragraphs = [
        (index, item, _normalize_for_match(item.text))
        for index, item in enumerate(paragraphs)
    ]
    matched_indices = set()
    matched_report_count = 0

    for segment in report_segments:
        segment_text = getattr(segment, "text", "") or ""
        normalized_segment = _normalize_for_match(segment_text)
        if len(normalized_segment) < 12:
            continue

        best_index = None
        best_score = 0.0
        for index, _, normalized_paragraph in normalized_paragraphs:
            if not normalized_paragraph:
                continue
            score = _partial_match_score(normalized_segment, normalized_paragraph)
            if score > best_score:
                best_score = score
                best_index = index

        if best_index is not None and best_score >= 0.68:
            matched_indices.add(best_index)
            matched_report_count += 1

    return [
        item
        for index, item in enumerate(paragraphs)
        if index in matched_indices
    ], matched_report_count


def _normalize_for_match(text: str) -> str:
    return re.sub(r"[\W_]+", "", (text or "").lower(), flags=re.UNICODE)


def _partial_match_score(needle: str, haystack: str) -> float:
    if not needle or not haystack:
        return 0.0
    if needle in haystack:
        return 1.0
    if haystack in needle and len(haystack) >= 20:
        return min(0.95, len(haystack) / len(needle))

    matcher = SequenceMatcher(None, needle, haystack, autojunk=False)
    match = matcher.find_longest_match(0, len(needle), 0, len(haystack))
    coverage = match.size / len(needle)
    if coverage >= 0.68:
        return coverage
    return matcher.ratio()


def _get_source_dir(token: str) -> Path:
    WORD_SOURCE_STORE_DIR.mkdir(parents=True, exist_ok=True)
    return WORD_SOURCE_STORE_DIR / token


def _load_word_source(filename: str, content: bytes) -> Tuple[DocxDocument, bytes]:
    suffix = Path(filename or "").suffix.lower()
    if suffix != ".docx":
        raise ValueError("仅支持上传 .docx 文件")

    return Document(io.BytesIO(content)), content


def _extract_body_paragraphs(document: DocxDocument) -> Tuple[List[ExtractedBodyParagraph], str]:
    extracted: List[ExtractedBodyParagraph] = []
    started = False
    in_toc = False
    collecting = False
    current_paragraph_type = "body"
    start_heading = ""

    for index, paragraph in enumerate(document.paragraphs):
        text = _normalize_paragraph_text(paragraph.text)
        style_name = paragraph.style.name if paragraph.style else ""
        if not text:
            continue

        if TOC_HEADING_RE.match(text) or _is_toc_style(style_name):
            in_toc = True
            collecting = False
            continue

        if STOP_SECTION_RE.match(text):
            if started:
                break
            continue

        if _is_keywords_heading(text):
            collecting = False
            continue

        if _is_abstract_heading(text) and not _is_toc_entry(text, style_name):
            collecting = False
            continue

        if _is_start_heading(text) and not _is_toc_entry(text, style_name):
            started = True
            in_toc = False
            collecting = True
            current_paragraph_type = "body"
            if not start_heading:
                start_heading = text
            continue

        if in_toc:
            continue

        if not started or not collecting:
            continue

        if _is_skippable_paragraph(text, style_name):
            continue

        extracted.append(
            ExtractedBodyParagraph(
                source_index=index,
                paragraph_type=current_paragraph_type,
                text=text,
            )
        )

    return extracted, start_heading


def _split_text_into_body_paragraphs(text: str) -> List[str]:
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    paragraphs = [part.strip() for part in re.split(r"\n\s*\n+", normalized) if part.strip()]
    return paragraphs


def _normalize_paragraph_text(text: str) -> str:
    return re.sub(r"[ \t\u3000]+", " ", text.replace("\r", " ").replace("\n", " ")).strip()


def _is_start_heading(text: str) -> bool:
    if FRONT_MATTER_HEADING_RE.match(text):
        return False
    return is_structured_heading_text(text)


def _is_abstract_heading(text: str) -> bool:
    return bool(ABSTRACT_HEADING_RE.match(text))


def _is_keywords_heading(text: str) -> bool:
    return bool(KEYWORDS_HEADING_RE.match(text))


def _is_skippable_paragraph(text: str, style_name: str) -> bool:
    style_lower = (style_name or "").lower()
    if PAGE_NUMBER_RE.match(text):
        return True
    if _is_toc_entry(text, style_name):
        return True
    if is_structured_heading_text(text):
        return True
    if SUMMARY_HEADING_RE.match(text):
        return True
    if CAPTION_RE.match(text):
        return True
    if any(token in style_lower for token in ("caption", "题注", "图题", "表题")):
        return True
    if any(token in style_lower for token in ("heading", "标题")) and _looks_like_style_only_heading(text):
        return True
    return False


def _is_toc_style(style_name: str) -> bool:
    style_lower = (style_name or "").lower()
    return "toc" in style_lower or "目录" in style_lower


def _is_toc_entry(text: str, style_name: str) -> bool:
    if _is_toc_style(style_name):
        return True
    if not re.search(r"\s+\d+$", text):
        return False

    title_part = re.sub(r"\s+\d+$", "", text).strip()
    if not title_part:
        return False

    if is_structured_heading_text(title_part):
        return True
    if SUMMARY_HEADING_RE.match(title_part):
        return True
    if STOP_SECTION_RE.match(title_part):
        return True
    return False


def is_structured_heading_text(text: str) -> bool:
    normalized = _normalize_paragraph_text(text)
    if not normalized:
        return False
    if MEASUREMENT_PREFIX_RE.match(normalized):
        return False

    if SINGLE_NUMERIC_HEADING_RE.match(normalized):
        tail = re.sub(r"^\d+\.(?!\d)\s*", "", normalized).strip()
        if not tail:
            return True
        if "：" in tail or ":" in tail:
            return False
        return len(tail) <= 24

    return any(
        pattern.match(normalized)
        for pattern in (
            FRONT_MATTER_HEADING_RE,
            START_HEADING_RE,
            CHAPTER_HEADING_RE,
            NUMERIC_HEADING_RE,
            COMPACT_NUMERIC_CHAPTER_HEADING_RE,
            CHINESE_SECTION_HEADING_RE,
            PAREN_CHINESE_HEADING_RE,
            INTRODUCTION_HEADING_RE,
            SUMMARY_HEADING_RE,
            STOP_SECTION_RE,
        )
    )


def _looks_like_style_only_heading(text: str) -> bool:
    normalized = _normalize_paragraph_text(text)
    if not normalized:
        return False
    if len(normalized) > 36:
        return False
    return not re.search(r"[。！？!?；;，,]", normalized)


def _apply_pale_yellow_highlight(paragraph) -> None:
    applied = False
    for run in paragraph.runs:
        if run.text:
            _set_shading(run._r.get_or_add_rPr())
            applied = True

    if not applied:
        _set_shading(paragraph._p.get_or_add_pPr())


def _set_shading(target) -> None:
    shading = target.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        target.append(shading)
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), PALE_YELLOW_FILL)


def _replace_paragraph_text_preserving_style(paragraph, text: str) -> None:
    first_run_props = None
    for run in paragraph.runs:
        if run._r.rPr is not None:
            first_run_props = deepcopy(run._r.rPr)
            break

    paragraph_element = paragraph._p
    for child in list(paragraph_element):
        if child.tag != qn("w:pPr"):
            paragraph_element.remove(child)

    if not text:
        return

    new_run = paragraph.add_run(text)
    if first_run_props is not None:
        run_props = new_run._r.get_or_add_rPr()
        for child in list(run_props):
            run_props.remove(child)
        for child in first_run_props:
            run_props.append(deepcopy(child))
